"""Export a chat summary (markdown text) to PDF / DOCX / Markdown bytes.

Used by the Streamlit chat to let a user save the agent's executive summary. The
summary text already contains the model figures (the agent is text-only); this
module just renders that text into a document with a non-advisory header — it adds
no numbers of its own. Pure (returns bytes); the heavy libs are lazy-imported.

PDF uses fpdf2's core latin-1 font, so unicode punctuation is transliterated to
ASCII for PDF only; DOCX and Markdown keep the text verbatim.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from io import BytesIO

DISCLAIMER = (
    "Research and education only - NOT financial advice. Figures come from statistical "
    "models and cited sources; the assistant provides no recommendations."
)

# fmt -> (mime type, file extension) for the download button.
EXPORT_META: dict[str, tuple[str, str]] = {
    "pdf": ("application/pdf", "pdf"),
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "md": ("text/markdown", "md"),
}

_UNICODE_MAP = {
    "→": "->", "←": "<-", "•": "-", "–": "-", "—": "--",
    "‘": "'", "’": "'", "“": '"', "”": '"', "…": "...",
    "−": "-", "×": "x", "✓": "[ok]", "✅": "[ok]", "❌": "[x]",
    "⚠": "[!]", "≈": "~", "≤": "<=", "≥": ">=", "±": "+/-",
}


def _ascii(s: str) -> str:
    """Transliterate common unicode to ASCII, then drop anything still non-latin-1."""
    for k, v in _UNICODE_MAP.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "ignore").decode("latin-1")


@dataclass
class _Block:
    """A parsed markdown block. ``kind`` in {h, bullet, p, blank, table}."""

    kind: str
    level: int = 0
    text: str = ""
    rows: list[list[str]] = field(default_factory=list)  # populated for tables


def _is_separator(line: str) -> bool:
    """A markdown table separator row, e.g. ``|---|:--:|``."""
    s = line.strip()
    return bool(s) and "-" in s and set(s) <= set("|-: ")


def _split_row(line: str) -> list[str]:
    """Split a ``| a | b |`` row into trimmed cells (bold markers stripped)."""
    return [c.strip().replace("**", "") for c in line.strip().strip("|").split("|")]


def _blocks(text: str) -> list[_Block]:
    """Parse markdown lines into blocks: headings, bullets, paragraphs, blanks, tables."""
    lines = text.splitlines()
    out: list[_Block] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        # Table: a row with pipes whose *next* line is a separator.
        if "|" in line and i + 1 < n and _is_separator(lines[i + 1]):
            header = _split_row(line)
            rows = [header]
            i += 2  # consume header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                row = _split_row(lines[i])
                ncol = len(header)
                rows.append((row + [""] * ncol)[:ncol])  # pad/truncate to header width
                i += 1
            out.append(_Block("table", rows=rows))
            continue
        if not line.strip():
            out.append(_Block("blank"))
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)", line)
        if h:
            out.append(_Block("h", level=len(h.group(1)), text=h.group(2).strip()))
            i += 1
            continue
        b = re.match(r"^\s*[-*+]\s+(.*)", line)
        if b:
            out.append(_Block("bullet", text=b.group(1).strip()))
            i += 1
            continue
        out.append(_Block("p", text=line.strip()))
        i += 1
    return out


def export_summary(
    text: str,
    fmt: str,
    *,
    title: str = "Stock Research Summary",
    generated_at: datetime | None = None,
    images: list[bytes] | None = None,
) -> bytes:
    """Render ``text`` (markdown) to ``fmt`` ('pdf' | 'docx' | 'md') as bytes.

    ``images`` are pre-rendered PNG bytes (chart figures) appended under a
    "Figures" heading in PDF/DOCX; Markdown is text-only (images are ignored).
    The exporter stays presentation-agnostic — it embeds bytes, it does not draw.
    """
    gen = generated_at or datetime.now(UTC)
    if fmt == "md":
        return _to_md(text, title, gen)
    if fmt == "docx":
        return _to_docx(text, title, gen, images)
    if fmt == "pdf":
        return _to_pdf(text, title, gen, images)
    raise ValueError(f"unsupported export format: {fmt!r} (use pdf/docx/md)")


def _stamp(gen: datetime) -> str:
    return f"Generated {gen:%Y-%m-%d %H:%M UTC}"


def _to_md(text: str, title: str, gen: datetime) -> bytes:
    header = f"# {title}\n\n*{_stamp(gen)}*\n\n> {DISCLAIMER}\n\n---\n\n"
    return (header + text.strip() + "\n").encode("utf-8")


def _add_runs(paragraph: object, content: str) -> None:
    """Add text to a docx paragraph, honoring **bold** spans."""
    for part in re.split(r"(\*\*[^*]+\*\*)", content):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])  # type: ignore[attr-defined]
            run.bold = True
        else:
            paragraph.add_run(part)  # type: ignore[attr-defined]


def _to_docx(text: str, title: str, gen: datetime, images: list[bytes] | None = None) -> bytes:
    from docx import Document  # lazy
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(_stamp(gen)).runs[0].italic = True
    doc.add_paragraph(DISCLAIMER)
    for blk in _blocks(text):
        if blk.kind == "h":
            doc.add_heading(blk.text.replace("**", ""), level=min(max(blk.level, 1), 4))
        elif blk.kind == "bullet":
            _add_runs(doc.add_paragraph(style="List Bullet"), blk.text)
        elif blk.kind == "p":
            _add_runs(doc.add_paragraph(), blk.text)
        elif blk.kind == "table" and blk.rows:
            ncol = len(blk.rows[0])
            table = doc.add_table(rows=0, cols=ncol)
            table.style = "Table Grid"
            for r_idx, row in enumerate(blk.rows):
                cells = table.add_row().cells
                for c_idx in range(ncol):
                    cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
                    if r_idx == 0:  # bold the header row
                        for para in cells[c_idx].paragraphs:
                            for run in para.runs:
                                run.bold = True
    if images:
        doc.add_heading("Figures", level=2)
        for png in images:
            doc.add_picture(BytesIO(png), width=Inches(6.0))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _to_pdf(text: str, title: str, gen: datetime, images: list[bytes] | None = None) -> bytes:
    from fpdf import FPDF  # lazy

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 9, _ascii(title), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(0, 6, _ascii(_stamp(gen)), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 8)
    pdf.multi_cell(0, 5, _ascii(DISCLAIMER), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    for blk in _blocks(text):
        if blk.kind == "blank":
            pdf.ln(2)
        elif blk.kind == "h":
            pdf.set_font("Helvetica", "B", max(11, 15 - blk.level))
            pdf.multi_cell(0, 7, _ascii(blk.text.replace("**", "")), new_x="LMARGIN", new_y="NEXT")
        elif blk.kind == "table" and blk.rows:
            _pdf_table(pdf, blk.rows)
        else:
            pdf.set_font("Helvetica", "", 11)
            body = ("- " + blk.text) if blk.kind == "bullet" else blk.text
            pdf.multi_cell(0, 6, _ascii(body), markdown=True, new_x="LMARGIN", new_y="NEXT")
    if images:
        pdf.set_font("Helvetica", "B", 13)
        pdf.ln(3)
        pdf.multi_cell(0, 7, "Figures", new_x="LMARGIN", new_y="NEXT")
        for png in images:
            pdf.ln(2)
            pdf.image(BytesIO(png), w=pdf.epw)  # scale to the content width; height auto
    return bytes(pdf.output())


def _pdf_table(pdf: object, rows: list[list[str]]) -> None:
    """Render a parsed markdown table with fpdf2's table API (first row = heading)."""
    pdf.set_font("Helvetica", "", 9)  # type: ignore[attr-defined]
    pdf.ln(1)  # type: ignore[attr-defined]
    with pdf.table(  # type: ignore[attr-defined]
        text_align="LEFT", line_height=5, padding=1.2, markdown=True
    ) as table:
        for row in rows:
            trow = table.row()
            for cell in row:
                trow.cell(_ascii(cell))
    pdf.ln(1)  # type: ignore[attr-defined]
